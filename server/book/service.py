from docx import Document
import io
from book.dao import BookDAO


async def generate_report(data: dict):
    # 1. Получаем данные
    raw_books = await BookDAO.find_all_full()

    # ПРЕОБРАЗОВАНИЕ: Если прилетели объекты ORM или Row,
    # переводим их в список чистых словарей для работы .get()
    books = []
    for b in raw_books:
        if hasattr(b, "_asdict"):  # Для Row (результаты селекта)
            books.append(b._asdict())
        elif hasattr(b, "__dict__"):  # Для обычных моделей
            books.append(b.__dict__)
        else:
            books.append(b)  # Если это уже словарь

    config = data.get('config', [])
    creator = data.get('creator', {})

    doc = Document()
    doc.add_heading('Сводный отчет по библиотеке', 0)

    # Составитель
    last_name = str(creator.get('last_name') or "")
    first_name = str(creator.get('first_name') or "")
    doc.add_paragraph(f"Составитель: {last_name} {first_name}")

    if not config:
        doc.add_paragraph("Конфигурация полей пуста.")
    else:
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=len(config))
        table.style = 'Table Grid'  # Можно оставить, если в шаблоне есть этот стиль

        # Заголовки
        for i, col in enumerate(config):
            header_text = str(col.get('field') or f"Колонка {i + 1}")
            table.rows[0].cells[i].text = header_text.capitalize()

        # Данные
        for book_dict in books:
            if not isinstance(book_dict, dict):
                continue  # Защита от пустых или битых записей

            row_cells = table.add_row().cells
            for i, col in enumerate(config):
                field_name = col.get('field')
                prefix = str(col.get('prefix') or "")
                suffix = str(col.get('suffix') or "")

                # Получаем значение
                raw_val = book_dict.get(field_name)

                # Обработка вложенных объектов (автор, жанр)
                if hasattr(raw_val, 'name'):  # Если это объект
                    val = str(raw_val.name)
                elif isinstance(raw_val, dict):  # Если это словарь
                    val = str(raw_val.get('name') or "—")
                elif raw_val is None:
                    val = "—"
                else:
                    val = str(raw_val)

                final_text = f"{prefix}{val}{suffix}".strip()
                row_cells[i].text = final_text if final_text else " "

    # Сохранение
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream